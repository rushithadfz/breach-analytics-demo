"""Multi-format text extraction (brief section 4).

Returns plain text with enough structural preservation (row/cell separators,
line breaks) that the regex/checksum detectors and the LLM tier both have
usable context. OCR-dependent formats raise OcrUnavailable if no OCR
backend is configured, rather than silently returning empty text — a
missing OCR engine is an infrastructure gap, not "no PII found".
"""
from __future__ import annotations

import csv
import email
import io
import re
import os
import shutil
from dataclasses import dataclass
from email.message import Message

import docx
import pdfplumber
from bs4 import BeautifulSoup
from openpyxl import load_workbook
from PIL import Image

from app.config import get_settings
from app.db.models import DocType

_TESSERACT_CHECKED = False
_TESSERACT_CMD: str | None = None

# Common Windows install location — winget/the official installer put it
# here and don't reliably add it to PATH, unlike Linux package managers.
_WINDOWS_FALLBACK_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


class OcrUnavailable(Exception):
    pass


def _tesseract_cmd() -> str | None:
    """Resolves the tesseract binary once: explicit config, then PATH,
    then the common Windows install location. Returns None if none of
    those actually exist."""
    global _TESSERACT_CHECKED, _TESSERACT_CMD
    if _TESSERACT_CHECKED:
        return _TESSERACT_CMD

    configured = get_settings().tesseract_cmd
    for candidate in (configured, shutil.which("tesseract"), _WINDOWS_FALLBACK_PATH):
        if candidate and os.path.isfile(candidate):
            _TESSERACT_CMD = candidate
            break
    _TESSERACT_CHECKED = True
    return _TESSERACT_CMD


def _pytesseract():
    cmd = _tesseract_cmd()
    if not cmd:
        raise OcrUnavailable("tesseract binary not found on PATH or at the configured/default location")
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = cmd
    return pytesseract


# --- Tiered OCR -------------------------------------------------------
#
# Measured (see design doc §1.2b — same images, same answer key):
#
#   engine           moderate scans   harsh scans   speed
#   Tesseract             100%            25%       0.73 s/doc
#   EasyOCR                --          1/6 fields   41.9 s/doc
#   GPT-5.5 vision        100%           100%       8.54 s/doc
#
# Neither extreme is defensible on its own. Paying vision-model cost for
# clean pages is waste when Tesseract is perfect on them and 12x faster;
# accepting 25% on degraded pages is indefensible when the output feeds a
# legal notification obligation.
#
# What makes escalation necessary rather than merely nice: classical OCR
# fails by NEAR-MISS, not by blank output. EasyOCR read a harsh scan's
# SSN as "234-04.6395" against a true "234-04-6395" — a period where a
# hyphen belongs. That is legible to a human and worthless to the
# pipeline: the value won't match, won't join to a person, and vanishes
# from the exposure table with nothing logged. So the escalation trigger
# cannot be "did OCR return text" — it has to be "does the text look
# like a successfully-read document".

# Escalation trigger: Tesseract's OWN mean per-word confidence.
#
# A first attempt used output length and a "gibberish ratio" instead, and
# it did not discriminate — a harsh scan that recovered only 25% of
# fields produced 126 characters against a 120-character floor, so it
# passed. Guessing at document shape was the wrong approach; the engine
# already publishes its uncertainty, and unlike a length heuristic it
# doesn't assume anything about how much text a document should contain.
#
# Measured on real pages (mean conf -> field accuracy):
#   real corpus scans   95.0, 95.4  -> 100%
#   benchmark moderate  83.0        -> 100%
#   benchmark harsh     64.0        ->  25%
#
# 75 sits in the gap. Everything that read correctly is comfortably
# above it; the read that lost three quarters of its fields is well
# below.
MIN_MEAN_OCR_CONFIDENCE = 75.0


def _tesseract_read_with_confidence(images: list) -> tuple[list[str], float]:
    """Returns (per-page texts, mean per-word confidence). One pass over
    the page — image_to_data gives both, so this costs no extra OCR work.

    Per-page rather than one joined string so the caller can record which
    page each character offset belongs to; joining is the caller's job."""
    pytesseract = _pytesseract()
    texts, confidences = [], []
    for img in images:
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        words = []
        for conf, word in zip(data["conf"], data["text"]):
            if not word.strip():
                continue
            words.append(word)
            c = str(conf)
            if c.lstrip("-").isdigit() and int(c) >= 0:
                confidences.append(int(c))
        texts.append(" ".join(words))
    mean_conf = sum(confidences) / len(confidences) if confidences else 0.0
    return texts, mean_conf


def _ocr_images_with_vision(images: list) -> list[str]:
    """Escalation path: GPT-5.5 reads the page directly. Uses the Azure
    deployment that already exists — creates no new service."""
    import base64
    import io as _io

    from openai import OpenAI

    from app.config import get_settings

    s = get_settings()
    if not s.azure_api_key:
        raise OcrUnavailable("vision escalation needs AZURE_API_KEY")
    client = OpenAI(api_key=s.azure_api_key, base_url=s.azure_openai_endpoint)

    out = []
    for img in images:
        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        resp = client.chat.completions.create(
            model=s.azure_strong_model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text":
                        "Transcribe every character of text in this scanned document exactly as "
                        "written. Preserve line breaks. Output only the transcription."},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                ],
            }],
            max_completion_tokens=2048,
        )
        out.append(resp.choices[0].message.content or "")
    return out


def _ocr_with_escalation(images: list, label: str) -> list[str]:
    """Tesseract first; escalate to vision only when Tesseract itself
    reports low confidence. Every fallback path returns Tesseract's text
    rather than raising — a degraded read is still worth more than a lost
    document, and the low confidence propagates to human review anyway.

    Returns one string per page."""
    pages, mean_conf = _tesseract_read_with_confidence(images)
    if mean_conf >= MIN_MEAN_OCR_CONFIDENCE:
        return pages

    try:
        from app.config import get_settings
        if not get_settings().azure_api_key:
            return pages  # no escalation configured
    except Exception:
        return pages

    try:
        vision = _ocr_images_with_vision(images)
        # Only accept the escalation if it kept the page structure — the
        # page map depends on a 1:1 correspondence with the rasterized
        # pages, and a short read that silently collapses them would
        # misattribute every offset after page 1.
        if len(vision) == len(pages):
            return vision
        return pages
    except Exception:
        return pages


def _page_images_from_pdf(path: str, dpi: int = 200) -> list:
    import pymupdf
    images = []
    with pymupdf.open(path) as doc:
        for page in doc:
            pix = page.get_pixmap(dpi=dpi)
            images.append(Image.frombytes("RGB", (pix.width, pix.height), pix.samples))
    return images


def _pages_image(path: str) -> list[str]:
    return _ocr_with_escalation([Image.open(path).convert("RGB")], path)


def _ocr_image(path: str) -> str:
    return "\n".join(_pages_image(path))


def _pages_pdf_digital(path: str) -> list[str]:
    with pdfplumber.open(path) as pdf:
        return [p.extract_text() or "" for p in pdf.pages]


def _extract_pdf_digital(path: str) -> str:
    return "\n".join(_pages_pdf_digital(path))


def _pages_pdf_scanned(path: str) -> list[str]:
    """Rasterizes each page with PyMuPDF (no external poppler dependency,
    unlike pdf2image), then runs the tiered OCR above."""
    return _ocr_with_escalation(_page_images_from_pdf(path), path)


def _extract_pdf_scanned(path: str) -> str:
    return "\n".join(_pages_pdf_scanned(path))


def _extract_docx(path: str) -> str:
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: str) -> str:
    # Read into memory rather than passing the path: openpyxl's path-based
    # loader rejects files whose extension isn't .xlsx, which breaks on the
    # wrong-extension edge case (real xlsx bytes saved with a .pdf name).
    with open(path, "rb") as f:
        buf = io.BytesIO(f.read())
    wb = load_workbook(buf, data_only=True)
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_csv(path: str) -> str:
    lines = []
    with open(path, encoding="utf-8") as f:
        for row in csv.reader(f):
            lines.append(" | ".join(row))
    return "\n".join(lines)


def _extract_html(path: str) -> str:
    with open(path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    # Keep table cell adjacency visible to regex context the same way the
    # generator wrote it (label, then value, as consecutive cells).
    return soup.get_text(separator=" | ")


def _extract_eml_body(path: str) -> str:
    with open(path, "rb") as f:
        msg: Message = email.message_from_binary_file(f)
    parts = [f"Subject: {msg.get('Subject', '')}", f"From: {msg.get('From', '')}"]
    for part in msg.walk():
        if part.get_content_disposition() == "attachment":
            continue
        if part.get_content_type() == "text/plain":
            payload = part.get_payload(decode=True)
            if payload:
                parts.append(payload.decode(part.get_content_charset() or "utf-8", errors="replace"))
    return "\n".join(parts)


def _extract_txt(path: str) -> str:
    with open(path, encoding="utf-8") as f:
        return f.read()


_DISPATCH = {
    DocType.pdf_digital: _extract_pdf_digital,
    DocType.pdf_scanned: _extract_pdf_scanned,
    DocType.docx: _extract_docx,
    DocType.xlsx: _extract_xlsx,
    DocType.csv: _extract_csv,
    DocType.html: _extract_html,
    DocType.eml: _extract_eml_body,
    DocType.txt: _extract_txt,
    DocType.png: _ocr_image,
}


def extract_text(doc_type: DocType, path: str) -> str:
    handler = _DISPATCH.get(doc_type)
    if handler is None:
        raise ValueError(f"no text extractor for {doc_type}")
    return handler(path)


def _looks_like_header(cells: list[str]) -> bool:
    """A header row is short, wordy, and carries no long digit runs."""
    if not cells:
        return False
    digits = sum(1 for c in cells if re.search(r"\d{4,}", c))
    return digits == 0 and all(len(c) <= 40 for c in cells)


def _label_row(headers: list[str] | None, cells: list[str]) -> str:
    """Re-attaches each cell to its column heading.

    A spreadsheet's meaning lives in its header row, and flattening a row
    to "a | b | c" throws that away — after which nothing downstream can
    tell an SSN cell from an order reference that happens to be shaped
    like one. That is not hypothetical: every row in this corpus carries
    an "Order Ref (not PII)" column deliberately formatted ###-##-####,
    and with the header discarded exactly half of all SSN values
    extracted from spreadsheets were that decoy — 377 fabricated SSNs
    attributed to real people.

    Emitting "Order Ref (not PII): 755-98-4598" instead of a bare cell
    restores the context the detectors already know how to use: the
    false-positive suppression sees the word "Ref", and the label-driven
    detectors see "Full Name" and "Home Address" as the labels they are,
    which is also why spreadsheet-only identities previously resolved
    with no name at all.

    Structure, not vocabulary: this works on any headed table without
    knowing what its columns are called.
    """
    if not headers:
        return " | ".join(cells)
    pairs = []
    for i, cell in enumerate(cells):
        head = headers[i].strip() if i < len(headers) else ""
        pairs.append(f"{head}: {cell}" if head else cell)
    return " | ".join(pairs)


def _rows_xlsx(path: str) -> list[tuple[str, str]]:
    with open(path, "rb") as f:
        buf = io.BytesIO(f.read())
    wb = load_workbook(buf, data_only=True)
    records = []
    for ws in wb.worksheets:
        headers: list[str] | None = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = ["" if c is None else str(c) for c in row]
            if not any(c.strip() for c in cells):
                continue
            if headers is None and _looks_like_header([c for c in cells if c]):
                headers = cells
                records.append((f"row:{i}", " | ".join(c for c in cells if c)))
                continue
            records.append((f"row:{i}", _label_row(headers, cells)))
    return records


def _rows_csv(path: str) -> list[tuple[str, str]]:
    records = []
    with open(path, encoding="utf-8") as f:
        headers: list[str] | None = None
        for i, row in enumerate(csv.reader(f)):
            if not row or not any(c.strip() for c in row):
                continue
            if headers is None and _looks_like_header([c for c in row if c]):
                headers = row
                records.append((f"row:{i}", " | ".join(c for c in row if c)))
                continue
            records.append((f"row:{i}", _label_row(headers, row)))
    return records


_ROW_DISPATCH = {DocType.xlsx: _rows_xlsx, DocType.csv: _rows_csv}


def extract_records(doc_type: DocType, path: str) -> list[tuple[str | None, str]]:
    """Returns [(record_key, text), ...]. Tabular formats split by row so
    entity resolution can tell two different people on two different rows
    of the same file apart; every other format is one record (record_key
    is None) covering the whole document."""
    handler = _ROW_DISPATCH.get(doc_type)
    if handler is not None:
        return handler(path)
    return [(None, extract_text(doc_type, path))]


# --- Page-aware parsing ----------------------------------------------
#
# Only some formats have a page at all. A PDF and a scanned image do; a
# CSV row, an email body and an HTML page do not, and inventing "page 1"
# for them would be a claim the format cannot support. Those keep an
# empty span list and their locator stays the record_key.

_PAGE_DISPATCH = {
    DocType.pdf_digital: _pages_pdf_digital,
    DocType.pdf_scanned: _pages_pdf_scanned,
    DocType.png: _pages_image,
}

PAGINATED_TYPES = frozenset(_PAGE_DISPATCH)


@dataclass
class ParsedRecord:
    record_key: str | None
    text: str
    # (page_number, start_offset, end_offset) over `text`, 1-indexed page
    # numbers, half-open offsets. Empty when the format has no pages.
    page_spans: list[tuple[int, int, int]]

    def page_for_offset(self, offset: int) -> int | None:
        """Which page a character offset falls on. Returns None when the
        format is unpaginated or the detector reported no position."""
        if offset < 0:
            return None
        for page_no, start, end in self.page_spans:
            if start <= offset < end:
                return page_no
        # An offset past the last span means the match landed on a
        # separator between pages; attribute it to the preceding page
        # rather than dropping the locator entirely.
        return self.page_spans[-1][0] if self.page_spans else None


def extract_parsed_records(doc_type: DocType, path: str) -> list[ParsedRecord]:
    """extract_records plus a page map, for the evidence deep-link.

    The joined text here MUST be assembled exactly the way extract_text
    assembles it ("\\n" between pages) — the page spans are offsets into
    that string, and any difference in the joiner would shift every page
    boundary and silently mis-cite the evidence."""
    pages_fn = _PAGE_DISPATCH.get(doc_type)
    if pages_fn is not None:
        pages = pages_fn(path)
        spans, cursor = [], 0
        for i, page_text in enumerate(pages, start=1):
            spans.append((i, cursor, cursor + len(page_text)))
            cursor += len(page_text) + 1  # +1 for the "\n" joiner
        return [ParsedRecord(None, "\n".join(pages), spans)]

    return [ParsedRecord(key, text, []) for key, text in extract_records(doc_type, path)]
